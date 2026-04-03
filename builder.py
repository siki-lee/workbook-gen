"""
练习册 Word 文档生成器 v2
严格按照「高端班题集样章326.docx」的排版规格还原
"""

import io
import os
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── 字体 ──────────────────────────────────────────────────────
FONT_SONG  = '宋体'
FONT_FANG  = '仿宋'
FONT_KAI   = '华文楷体'
FONT_KAI2  = '楷体'   # 描红用
FONT_TIMES = 'Times New Roman'

# ── 颜色 ──────────────────────────────────────────────────────
# 核心方法
C_CM_TITLE   = '984807'   # 棕色  标题字
C_CM_BORDER  = '2D5A8E'   # 蓝色  下边框
C_CM_LINE    = 'D1D5DB'   # 浅灰  空白行下边框

# 部分标题
C_SEC_BG     = 'EFF6FF'   # 浅蓝背景
C_SEC_LINE   = 'BFDBFE'   # 蓝底边
C_SEC_NUM    = '2D5A8E'   # 第X部分 字色
C_SEC_NAME   = '1A3A5C'   # 部分名称字色
C_SEC_TIME   = '6B7280'   # 建议用时字色

# 特殊版块标题（妙笔生花/文心雕龙/日积月累）
C_SPECIAL    = '2D5A8E'

# 题型标签
TYPE_COLORS = {
    '核心题': 'E8722A',
    '热搜题': '4A9B8E',
    '新趋势': '7DB640',
}

# 描红
C_MIAOHONG_TEXT  = 'C0C0C0'  # 灰色字
C_MIAOHONG_BG    = 'FAFAFA'  # 浅白背景
C_TABLE_BORDER   = 'BBBBBB'
C_WRITE_DASHED   = 'CCCCCC'


# ══════════════════════════════════════════════════════════════
# 底层工具
# ══════════════════════════════════════════════════════════════

def _set_cn_font(run, name):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), name)


def _set_shading_run(run, fill_hex):
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    rPr.append(shd)


def _set_shading_para(para, fill_hex):
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    pPr.append(shd)


def _set_para_border(para, bottom=None, top=None, left=None, right=None):
    """bottom/top/left/right: (color, sz, space, val) 或 None"""
    pPr = para._element.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)

    def _add(tag, color, sz, space, val='single'):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:val'),   val)
        el.set(qn('w:color'), color)
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), str(space))
        pBdr.append(el)

    if bottom: _add('bottom', *bottom)
    if top:    _add('top',    *top)
    if left:   _add('left',   *left)
    if right:  _add('right',  *right)


def _spacing(para, before_pt=0, after_pt=0, line_twips=None, line_rule=None):
    fmt = para.paragraph_format
    if before_pt:
        fmt.space_before = Pt(before_pt)
    if after_pt:
        fmt.space_after = Pt(after_pt)
    if line_twips:
        pPr = para._element.get_or_add_pPr()
        sp  = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        sp.set(qn('w:line'),     str(line_twips))
        sp.set(qn('w:lineRule'), line_rule or 'auto')


def _run(para, text, font=FONT_SONG, size=None, bold=False,
         color=None, shd=None, italic=False, align_center=False):
    if align_center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(text)
    _set_cn_font(r, font)
    if size:   r.font.size = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:  r.font.color.rgb = RGBColor.from_string(color)
    if shd:    _set_shading_run(r, shd)
    return r


def _page_break(doc):
    p  = doc.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p.add_run()._element.append(br)


# ══════════════════════════════════════════════════════════════
# 核心方法
# ══════════════════════════════════════════════════════════════

def _add_core_method_area(doc, blank_lines=8):
    """
    核心方法标题（棕色，蓝色下边框）
    + 空白书写区（浅灰横线，供学生自主复习）
    """
    # 标题行
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before_pt=8, after_pt=4)
    _set_para_border(p, bottom=(C_CM_BORDER, 4, 4))
    _run(p, '核心方法', font=FONT_SONG, size=13, bold=True, color=C_CM_TITLE)

    # 空白书写行（浅灰底边线）
    for _ in range(blank_lines):
        lp = doc.add_paragraph()
        _spacing(lp, before_pt=1.5, after_pt=1.5)
        _set_para_border(lp, bottom=(C_CM_LINE, 2, 2))


# ══════════════════════════════════════════════════════════════
# 版块标题
# ══════════════════════════════════════════════════════════════

SECTION_LABELS = ['第一部分', '第二部分', '第三部分', '第四部分']


def _add_section_header(doc, idx, name, time_suggestion=''):
    """
    第X部分  巩固提升/能力进阶
    蓝色背景 + 蓝色底边框，深蓝文字
    """
    label = SECTION_LABELS[idx] if idx < len(SECTION_LABELS) else f'第{idx+1}部分'
    p = doc.add_paragraph()
    _spacing(p, before_pt=5, after_pt=3)
    _set_shading_para(p, C_SEC_BG)
    _set_para_border(p, bottom=(C_SEC_LINE, 2, 2))

    _run(p, label,           font=FONT_SONG, size=14, bold=True, color=C_SEC_NUM)
    _run(p, f'    {name}',   font=FONT_SONG, size=14, bold=True, color=C_SEC_NAME)
    if time_suggestion:
        _run(p, f'    建议用时 {time_suggestion}', font=FONT_SONG, size=9, color=C_SEC_TIME)


def _add_special_header(doc, title):
    """
    妙笔生花 / 文心雕龙 / 日积月累 标题
    深蓝色，14pt，无背景
    """
    p = doc.add_paragraph()
    _spacing(p, before_pt=8, after_pt=4)
    _run(p, title, font=FONT_SONG, size=14, bold=True, color=C_SPECIAL)


# ══════════════════════════════════════════════════════════════
# 阅读文章
# ══════════════════════════════════════════════════════════════

def _add_article(doc, title, body, source='', image_bytes=None, author=''):
    # 文章标题
    if title:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(tp, before_pt=6, after_pt=2)
        _run(tp, title, font=FONT_SONG, size=13, bold=True)

    # 作者
    if author and author.strip():
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(ap, before_pt=0, after_pt=3)
        _run(ap, author.strip(), font=FONT_FANG, size=11)

    # 正文
    if body:
        for line in body.strip().split('\n'):
            lp = doc.add_paragraph()
            lp.paragraph_format.first_line_indent = Pt(21)  # 首行缩进两字符（10.5pt × 2）
            _spacing(lp, before_pt=0, after_pt=0, line_twips=288, line_rule='auto')
            _run(lp, line.strip(), font=FONT_KAI, size=10.5)

    # 图片
    if image_bytes:
        _insert_image(doc, image_bytes, width_cm=14)

    # 出处
    if source and source.strip():
        sp = doc.add_paragraph()
        _spacing(sp, before_pt=2, after_pt=2)
        sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(sp, source.strip(), font=FONT_SONG, size=10.5)


# ══════════════════════════════════════════════════════════════
# 题目（阅读类）
# ══════════════════════════════════════════════════════════════

def _add_question_table(doc, table_data):
    """在题目中渲染表格，首行可选表头样式"""
    data = table_data.get('data', [])
    has_header = table_data.get('has_header', True)
    if not data or not data[0]:
        return

    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    for r_i, row in enumerate(data):
        is_header = has_header and r_i == 0
        for c_i, cell_text in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before_pt=2, after_pt=2)
            r = p.add_run(cell_text)
            _set_cn_font(r, FONT_SONG)
            r.font.size = Pt(10.5)
            r.font.bold = is_header
            if is_header:
                tcPr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'EFF6FF')
                tcPr.append(shd)

    doc.add_paragraph()  # 表格后留空


def _add_question(doc, number, type_tag, text,
                  linked_material='', hint='', answer_lines=4,
                  image_bytes=None, table_data=None):
    fill = TYPE_COLORS.get(type_tag, '888888')

    # 题号 + 标签 + 题目文字
    qp = doc.add_paragraph()
    _spacing(qp, before_pt=8, after_pt=3)
    _run(qp, f'{number} ',  font=FONT_SONG, size=12, bold=True,  color='333333')
    _run(qp, f' {type_tag} ', font=FONT_SONG, size=9,  bold=True, color='FFFFFF', shd=fill)
    _run(qp, f'   {text}', font=FONT_SONG, size=11)

    # 表格（题目附表）
    if table_data:
        _add_question_table(doc, table_data)

    # 题目图片
    if image_bytes:
        _insert_image(doc, image_bytes, width_cm=14)

    # 链接材料
    if linked_material and linked_material.strip():
        lmp = doc.add_paragraph()
        _spacing(lmp, before_pt=2, after_pt=2)
        _run(lmp, '【链接材料】', font=FONT_SONG, size=11)
        for line in linked_material.strip().split('\n'):
            lp = doc.add_paragraph()
            _spacing(lp, before_pt=0, after_pt=0)
            _run(lp, line.strip(), font=FONT_KAI, size=11)

    # 答题提示
    if hint and hint.strip():
        hp = doc.add_paragraph()
        _spacing(hp, before_pt=2, after_pt=2)
        _run(hp, f'答题提示：{hint.strip()}', font=FONT_FANG, size=10.5)

    # 答题空白行
    for _ in range(answer_lines):
        ap = doc.add_paragraph()
        _spacing(ap, before_pt=0, after_pt=0, line_twips=480, line_rule='exact')


# ══════════════════════════════════════════════════════════════
# 作文专用
# ══════════════════════════════════════════════════════════════

def _add_writing_prompt(doc, prompt_text, requirements='', writing_lines=22):
    """作文题目 + 要求 + 空白书写行"""
    # 题目
    if prompt_text:
        pp = doc.add_paragraph()
        _spacing(pp, before_pt=2, after_pt=2)
        _run(pp, prompt_text, font=FONT_SONG, size=11)

    # 要求
    if requirements:
        rp = doc.add_paragraph()
        _spacing(rp, before_pt=2, after_pt=4)
        _run(rp, requirements, font=FONT_SONG, size=11)

    # 空白书写行（浅灰底线）
    for _ in range(writing_lines):
        wp = doc.add_paragraph()
        _spacing(wp, before_pt=1, after_pt=1, line_twips=480, line_rule='exact')
        _set_para_border(wp, bottom=(C_CM_LINE, 2, 2))


def _add_miaobishenghua(doc, writing_lines=25):
    """妙笔生花：说明文字 + 大量空白书写行"""
    _add_special_header(doc, '妙笔生花')

    tip = doc.add_paragraph()
    _spacing(tip, before_pt=2, after_pt=4)
    _run(tip, '请在上面的题目中，挑选一个你最心动的，尽情挥洒你的笔墨吧！',
         font=FONT_SONG, size=10)

    for _ in range(writing_lines):
        wp = doc.add_paragraph()
        _spacing(wp, before_pt=1, after_pt=1, line_twips=480, line_rule='exact')
        _set_para_border(wp, bottom=(C_CM_LINE, 2, 2))


def _add_wenxindiaolong(doc, criteria=None):
    """文心雕龙：自评打分表"""
    _add_special_header(doc, '文心雕龙')

    tip = doc.add_paragraph()
    _spacing(tip, before_pt=2, after_pt=4)
    _run(tip, '恭喜你完成本讲习作，快来自评一下！如果你满足了标准，就在对应的框里打上"√"吧！',
         font=FONT_SONG, size=10)

    if criteria is None:
        criteria = [
            ('立意与审题', '立意深刻，扣题紧，有普遍意义'),
            ('内容与选材', '选材独特，细节真实生动，有感染力'),
            ('结构与逻辑', '结构清晰，情感有层次，收尾呼应'),
            ('语言与表达', '语言有个性，善用修辞，文笔流畅'),
        ]

    table = doc.add_table(rows=len(criteria) + 2, cols=5)
    table.style = 'Table Grid'

    # 表头行1
    header1 = table.rows[0]
    _set_table_cell(header1.cells[0], '评分维度', bold=True, bg='EFF6FF')
    _set_table_cell(header1.cells[1], '评价标准', bold=True, bg='EFF6FF')
    header1.cells[2].merge(header1.cells[3]).merge(header1.cells[4])
    _set_table_cell(header1.cells[2], '等级', bold=True, bg='EFF6FF', center=True)

    # 表头行2
    header2 = table.rows[1]
    _set_table_cell(header2.cells[0], '', bold=True, bg='EFF6FF')
    _set_table_cell(header2.cells[1], '', bold=True, bg='EFF6FF')
    for j, grade in enumerate(['优秀', '良好', '一般']):
        _set_table_cell(header2.cells[2 + j], grade, bold=True, bg='EFF6FF', center=True)

    # 内容行
    for i, (dim, std) in enumerate(criteria):
        row = table.rows[i + 2]
        _set_table_cell(row.cells[0], dim,  bold=True)
        _set_table_cell(row.cells[1], std)
        for j in range(3):
            _set_table_cell(row.cells[2 + j], '', center=True)


def _set_table_cell(cell, text, bold=False, bg=None, center=False, size=10):
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before_pt=2, after_pt=2)
    r = p.add_run(text)
    _set_cn_font(r, FONT_SONG)
    r.font.size = Pt(size)
    r.font.bold = bold
    if bg:
        tcPr = cell._element.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  bg)
        tcPr.append(shd)


# ══════════════════════════════════════════════════════════════
# 日积月累（描红表格）
# ══════════════════════════════════════════════════════════════

def _add_rijiyuelei(doc, words):
    """
    日积月累版块
    words: list of {'pinyin': 'lǎng rùn', 'hanzi': '朗润'}
    每5个词一组，生成3行：拼音行、描红行、默写行
    """
    _add_special_header(doc, '日积月累')

    tip = doc.add_paragraph()
    _spacing(tip, before_pt=2, after_pt=6)
    _run(tip, '先在灰色字上描红记忆字形，再在下方空格中独立默写。',
         font=FONT_SONG, size=10, color='666666')

    # 按每行5个分组
    chunk_size = 5
    groups = [words[i:i + chunk_size] for i in range(0, len(words), chunk_size)]
    total_words = len(words)
    n = 0  # 全局序号

    for group in groups:
        cols = chunk_size
        # 3行：拼音行、描红行、默写行
        table = doc.add_table(rows=3, cols=cols)

        # 行0：拼音
        for j, w in enumerate(group):
            cell = table.rows[0].cells[j]
            _fill_cell_no_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before_pt=0, after_pt=0)
            r = p.add_run(f'{n + j + 1}. {w["pinyin"]}')
            r.font.name      = FONT_TIMES
            r.font.size      = Pt(10)
            r.font.color.rgb = RGBColor(0, 0, 0)

        # 行1：灰色描红汉字
        for j, w in enumerate(group):
            cell = table.rows[1].cells[j]
            _fill_cell_with_style(cell, bg=C_MIAOHONG_BG, border_color=C_TABLE_BORDER)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before_pt=2, after_pt=2)
            r = p.add_run(w['hanzi'])
            _set_cn_font(r, FONT_KAI2)
            r.font.size      = Pt(18)
            r.font.color.rgb = RGBColor.from_string(C_MIAOHONG_TEXT)

        # 行2：空白默写格
        for j in range(cols):
            cell = table.rows[2].cells[j]
            _fill_cell_write(cell)
            p = cell.paragraphs[0]
            _spacing(p, before_pt=8, after_pt=8)

        n += len(group)

        # 组间空隙
        doc.add_paragraph()


def _fill_cell_no_border(cell):
    tcPr = cell._element.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _fill_cell_with_style(cell, bg, border_color):
    tcPr = cell._element.get_or_add_tcPr()
    # 背景
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg)
    tcPr.append(shd)
    # 边框
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:color'), border_color)
        el.set(qn('w:sz'),    '4')
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _fill_cell_write(cell):
    """默写行：上边虚线，其余实线"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),   'dashed')
    top.set(qn('w:color'), C_WRITE_DASHED)
    top.set(qn('w:sz'),    '4')
    tcBdr.append(top)
    for side in ('left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:color'), C_TABLE_BORDER)
        el.set(qn('w:sz'),    '4')
        tcBdr.append(el)
    tcPr.append(tcBdr)


# ══════════════════════════════════════════════════════════════
# 图片插入
# ══════════════════════════════════════════════════════════════

def _add_answers_section(doc, lectures):
    """
    参考答案汇总版块，排在文档最末尾。
    只渲染有答案内容的题目。
    """
    # 收集所有有答案的题目
    entries = []
    for lecture in lectures:
        subject = lecture.get('subject', '')
        number  = lecture.get('number', '')
        topic   = lecture.get('topic', '')
        sections = lecture.get('sections', [])
        lec_entries = []
        for sec_idx, section in enumerate(sections[:2]):
            sec_name = section.get('name', '')
            sec_entries = []
            for q_idx, q in enumerate(section.get('questions', []), start=1):
                answer = q.get('answer', '').strip()
                if answer:
                    sec_entries.append((q_idx, q.get('text', ''), answer))
            if sec_entries:
                lec_entries.append((sec_idx, sec_name, sec_entries))
        if lec_entries:
            entries.append((number, subject, topic, lec_entries))

    if not entries:
        return

    _page_break(doc)

    # 大标题
    p = doc.add_paragraph()
    _spacing(p, before_pt=8, after_pt=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, '参考答案', font=FONT_SONG, size=16, bold=True, color=C_SPECIAL)

    for number, subject, topic, lec_entries in entries:
        # 讲次标题
        lp = doc.add_paragraph()
        _spacing(lp, before_pt=10, after_pt=4)
        _set_shading_para(lp, C_SEC_BG)
        _set_para_border(lp, bottom=(C_SEC_LINE, 2, 2))
        label = f'第 {number} 讲'
        if subject:
            label += f'  【{subject}】'
        if topic:
            label += f'  {topic}'
        _run(lp, label, font=FONT_SONG, size=12, bold=True, color=C_SEC_NAME)

        for sec_idx, sec_name, sec_entries in lec_entries:
            # 部分标题
            sec_label = SECTION_LABELS[sec_idx] if sec_idx < len(SECTION_LABELS) else f'第{sec_idx+1}部分'
            sp = doc.add_paragraph()
            _spacing(sp, before_pt=6, after_pt=2)
            _run(sp, f'{sec_label}  {sec_name}', font=FONT_SONG, size=11, bold=True, color=C_SEC_NUM)

            for q_idx, q_text, answer in sec_entries:
                # 题号 + 答案
                ap = doc.add_paragraph()
                _spacing(ap, before_pt=3, after_pt=2)
                _run(ap, f'第{q_idx}题　', font=FONT_SONG, size=10.5, bold=True, color='333333')
                _run(ap, answer, font=FONT_SONG, size=10.5)


def _insert_image(doc, image_bytes, width_cm=14):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    tmp.write(image_bytes)
    tmp.close()
    try:
        ip  = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(tmp.name, width=Cm(width_cm))
    finally:
        os.unlink(tmp.name)


# ══════════════════════════════════════════════════════════════
# 主入口
# ══════════════════════════════════════════════════════════════

# 板块类型常量
SUBJECT_READING  = '现代文阅读'
SUBJECT_WRITING  = '作文'
SUBJECT_BASIC    = '基础'
SUBJECT_CLASSICS = '名著阅读'
SUBJECT_WENYAN   = '文言文阅读'
SUBJECT_POETRY   = '诗词'


def build_document(data: dict) -> io.BytesIO:
    """
    data 结构见 app.py 注释。
    板块类型决定生成哪些 section：
      现代文阅读 / 名著阅读 / 文言文阅读 / 诗词:
          核心方法 + 巩固提升 + 能力进阶
      作文:
          核心方法 + 巩固提升 + 能力进阶 + 妙笔生花 + 文心雕龙
      基础:
          核心方法 + 巩固提升 + 能力进阶 + 日积月累
    """
    doc = Document()

    # 页面设置
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(1.91)
    sec.bottom_margin = Cm(1.91)
    sec.left_margin   = Cm(1.91)
    sec.right_margin  = Cm(1.91)

    # 清除默认段落
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    for lec_idx, lecture in enumerate(data.get('lectures', [])):
        if lec_idx > 0:
            _page_break(doc)

        subject = lecture.get('subject', SUBJECT_READING)
        number  = lecture.get('number', lec_idx + 1)
        topic   = lecture.get('topic', '')

        # ── 讲次标题 ──
        p1 = doc.add_paragraph()
        _spacing(p1, before_pt=8, after_pt=3)
        _run(p1, f'第 {number} 讲 ', font=FONT_SONG, size=10.5, bold=True)

        p2 = doc.add_paragraph()
        _spacing(p2, before_pt=4, after_pt=4)
        if subject:
            _run(p2, f'【{subject}】', font=FONT_SONG, size=22, bold=True)
        _run(p2, topic, font=FONT_SONG, size=22, bold=True)

        # ── 核心方法（始终是空白复习区）──
        blank = lecture.get('core_method_lines', 8)
        _add_core_method_area(doc, blank_lines=blank)

        # ── 根据板块类型生成内容 ──
        sections = lecture.get('sections', [])  # 巩固提升、能力进阶

        if subject == SUBJECT_WRITING:
            # 作文：巩固提升 + 能力进阶 = 写作题目 + 空行
            for sec_idx, section in enumerate(sections[:2]):
                _add_section_header(doc, sec_idx, section.get('name', ''), section.get('time_suggestion', ''))
                for prompt in section.get('prompts', []):
                    _add_writing_prompt(
                        doc,
                        prompt.get('text', ''),
                        prompt.get('requirements', ''),
                        prompt.get('writing_lines', 22),
                    )
            # 妙笔生花 + 文心雕龙
            _add_miaobishenghua(doc)
            _add_wenxindiaolong(doc)

        elif subject == SUBJECT_BASIC:
            # 基础：巩固提升 + 能力进阶 = 阅读/题目；最后加日积月累
            for sec_idx, section in enumerate(sections[:2]):
                _add_section_header(doc, sec_idx, section.get('name', ''), section.get('time_suggestion', ''))
                for article in section.get('articles', []):
                    _add_article(doc, article.get('title', ''), article.get('body', ''),
                                 article.get('source', ''), article.get('image'),
                                 article.get('author', ''))
                for q_idx, q in enumerate(section.get('questions', []), start=1):
                    _add_question(doc, q_idx, q.get('type', '核心题'), q.get('text', ''),
                                  q.get('linked_material', ''), q.get('hint', ''),
                                  q.get('answer_lines', 4), q.get('image'),
                                  q.get('table'))
            # 日积月累
            words = lecture.get('daily_words', [])
            if words:
                _add_rijiyuelei(doc, words)

        else:
            # 现代文阅读 / 名著 / 文言文 / 诗词：阅读文章 + 题目
            for sec_idx, section in enumerate(sections[:2]):
                _add_section_header(doc, sec_idx, section.get('name', ''), section.get('time_suggestion', ''))
                for article in section.get('articles', []):
                    _add_article(doc, article.get('title', ''), article.get('body', ''),
                                 article.get('source', ''), article.get('image'),
                                 article.get('author', ''))
                for q_idx, q in enumerate(section.get('questions', []), start=1):
                    _add_question(doc, q_idx, q.get('type', '核心题'), q.get('text', ''),
                                  q.get('linked_material', ''), q.get('hint', ''),
                                  q.get('answer_lines', 4), q.get('image'),
                                  q.get('table'))

    _add_answers_section(doc, data.get('lectures', []))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
